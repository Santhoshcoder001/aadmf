from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_step(doc, number, title, action, why, checks=None):
    add_heading(doc, f"Step {number}: {title}", level=2)
    add_para(doc, "What to do:", bold=True)
    doc.add_paragraph(action, style="List Bullet")
    add_para(doc, "Why this step is done:", bold=True)
    doc.add_paragraph(why)
    if checks:
        add_para(doc, "What to verify:", bold=True)
        for c in checks:
            doc.add_paragraph(c, style="List Bullet")


def build_document(output_path: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "AADMF - Full Project Process and Execution Guide", level=0)
    doc.add_paragraph("Date: April 6, 2026")

    add_heading(doc, "1. Project Explanation", level=1)
    doc.add_paragraph(
        "Adaptive Agentic Drift Mining Framework (AADMF) is a multi-agent Python project designed for streaming data mining under distribution drift. "
        "It combines drift detection, algorithm planning, adaptive mining, hypothesis generation and validation, and tamper-evident provenance logging in one workflow."
    )

    add_heading(doc, "2. What the Project Does", level=1)
    for item in [
        "Detects drift on incoming batches using a Page-Hinkley style detector.",
        "Selects the most suitable mining algorithm per batch using a Planner Agent.",
        "Runs miners such as IsolationForest, DBSCAN, or StatisticalRules.",
        "Generates and validates hypotheses when drift is meaningful.",
        "Stores provenance events as a hash-chained audit trail and can optionally use Neo4j.",
        "Visualizes run outputs through a Streamlit dashboard.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "3. Core Components", level=1)
    components = [
        ("aadmf/agents", "Planner, Miner, Hypothesizer, and Validator agent logic."),
        ("aadmf/drift", "Drift detector implementation."),
        ("aadmf/orchestrator", "Manual and LangGraph orchestrators that control flow."),
        ("aadmf/provenance", "Hash-chain logger and optional Neo4j-backed logger."),
        ("aadmf/dashboard", "Streamlit dashboard for metrics and provenance visualization."),
        ("poc.py", "Main script to execute end-to-end pipeline."),
        ("config.yaml", "Configuration for execution mode, data, LLM, and provenance."),
    ]
    for name, desc in components:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    add_heading(doc, "4. Full Step-by-Step Execution Process", level=1)

    add_step(
        doc,
        1,
        "Open the project folder",
        "In PowerShell, run: cd C:/Users/santh/OneDrive/Desktop/aadmf/aadmf-main",
        "All relative paths in scripts and the dashboard expect the repository root as current working directory.",
    )

    add_step(
        doc,
        2,
        "Create and activate virtual environment",
        "Run: python -m venv .venv\nThen: ./.venv/Scripts/Activate.ps1\nThen: python -m pip install --upgrade pip",
        "This isolates dependencies for the project and avoids package conflicts with system Python.",
    )

    add_step(
        doc,
        3,
        "Install project dependencies",
        "Run: pip install pyyaml pandas numpy scipy scikit-learn langgraph ucimlrepo neo4j streamlit plotly pyvis streamlit-autorefresh ollama pytest",
        "The project requires these libraries for data loading, model execution, orchestration, provenance, UI, and testing.",
    )

    add_step(
        doc,
        4,
        "Configure data source in config.yaml",
        "Set uci_loader and uci_streaming data_dir values to your batch folder (for this repo, dataset1 is common).",
        "The pipeline must know where batch files are located before it can process data.",
    )

    add_step(
        doc,
        5,
        "Set execution mode",
        "In config.yaml, set execution.mode to full and execution.use_langgraph to true (recommended).",
        "This controls orchestrator behavior and ensures graph-based routing is exercised.",
    )

    add_step(
        doc,
        6,
        "Run the end-to-end pipeline",
        "Run: python poc.py",
        "This is the primary run that executes drift detection, planning, mining, hypothesis logic, and provenance logging.",
        checks=[
            "Batch-wise output appears in terminal.",
            "provenance.json is created or updated.",
            "Selected algorithms and drift scores are visible.",
        ],
    )

    add_step(
        doc,
        7,
        "Run UCI loader smoke test",
        "Run: python test_uci_loader.py",
        "Confirms that the configured loader and basic pipeline path are working correctly before deeper tests.",
    )

    add_step(
        doc,
        8,
        "Run LangGraph parity and routing test",
        "Run: python test_langgraph.py",
        "Validates parity between orchestrators and confirms conditional branch behavior under low/high drift.",
    )

    add_step(
        doc,
        9,
        "Run unit tests",
        "Run: python -m pytest tests/unit/test_planner.py -q",
        "Checks Planner Agent scoring and update behavior, which directly impacts algorithm selection quality.",
    )

    add_step(
        doc,
        10,
        "Launch Streamlit dashboard",
        "Run: streamlit run aadmf/dashboard/app.py",
        "Dashboard is used to inspect run quality, drift trend, algorithm distribution, hypotheses, and provenance integrity.",
    )

    add_step(
        doc,
        11,
        "Verify dashboard sidebar paths",
        "Ensure provenance path is provenance.json and config path is config.yaml.",
        "If incorrect paths are used, dashboard may appear empty even when pipeline output exists.",
    )

    add_step(
        doc,
        12,
        "Validate dashboard outputs",
        "Check all sections: drift chart, algorithm chart, hypothesis feed, provenance graph, and tamper interface.",
        "This confirms the system produced coherent outputs and that visualization receives the expected schema.",
    )

    add_step(
        doc,
        13,
        "Run tamper demonstration",
        "In dashboard, click Run tamper demo.",
        "Demonstrates that hash-chain provenance detects modifications and supports audit integrity.",
    )

    add_step(
        doc,
        14,
        "Optional: enable Ollama hypothesis phrasing",
        "Run: ollama pull phi3:mini and python test_llm_hypothesis.py\nSet hypothesizer.use_llm=true and llm.model=phi3:mini in config.yaml",
        "This enables LLM-based statement phrasing while preserving fallback behavior.",
    )

    add_step(
        doc,
        15,
        "Optional: enable Neo4j provenance backend",
        "Set provenance.backend=neo4j with valid Neo4j credentials in config.yaml, then run python poc.py again.",
        "Allows event persistence and graph-native provenance queries beyond local JSON export.",
    )

    add_step(
        doc,
        16,
        "Optional: run week-6 tuning workflow",
        "Run: python tune_week6.py",
        "Produces experiment outputs for tuning and recommendation artifacts.",
        checks=[
            "experiments/results/scoring_matrix_tuning.csv",
            "experiments/results/miner_tuning_results.csv",
            "experiments/results/week6_final_recommendation.md",
        ],
    )

    add_heading(doc, "5. Quick End-to-End Command Block", level=1)
    quick_cmds = [
        "python -m venv .venv",
        "./.venv/Scripts/Activate.ps1",
        "python -m pip install --upgrade pip",
        "pip install pyyaml pandas numpy scipy scikit-learn langgraph ucimlrepo neo4j streamlit plotly pyvis streamlit-autorefresh ollama pytest",
        "python poc.py",
        "python test_uci_loader.py",
        "python test_langgraph.py",
        "python -m pytest tests/unit/test_planner.py -q",
        "streamlit run aadmf/dashboard/app.py",
    ]
    for cmd in quick_cmds:
        doc.add_paragraph(cmd, style="List Bullet")

    add_heading(doc, "6. Why this process order is recommended", level=1)
    for point in [
        "Setup and dependency steps are done first to prevent runtime import failures.",
        "Configuration is done before execution so the pipeline can find data and choose the intended orchestration mode.",
        "Main pipeline run is done before dashboard so visualizations have real output to display.",
        "Smoke and parity tests are run before deep experimentation to catch integration issues early.",
        "Optional features (Ollama, Neo4j, tuning) are last because they are not required for baseline execution.",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    doc.save(output_path)


if __name__ == "__main__":
    build_document("AADMF_Full_Project_Process.docx")
    print("Created AADMF_Full_Project_Process.docx")
